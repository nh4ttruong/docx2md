import docx
from pathlib import Path
import os, glob

#delete three first line in docx file
def delete_lines(dires, line_num=3):
    #get all docx file in folder
    os.chdir(dires)
    paths = glob.glob('*.docx')

    for path in paths:
        if os.path.basename(path).endswith('.docx'):
            doc = docx.Document(path)
            for i in range(line_num):
                doc.paragraphs[i].clear()
                #delete empty line
            if doc.paragraphs[0].text == '' or doc.paragraphs[1].text == '':
                doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._p)
                doc.paragraphs[1]._element.getparent().remove(doc.paragraphs[1]._p)
            
            p = dires + new_folder + '\\' #save file to edited_docx folder
            os.makedirs(name=p, exist_ok=True)
            
            #rename file
            new_name = name_file(path)
            
            #save file to edited_docx folder
            doc.save(p + new_name)

def docx_2_md(dires):
    #new dires
    dires = dires + new_folder + "\\"
    #get all docx file in folder
    os.chdir(dires)
    paths = glob.glob('*.docx')
    print(dires + os.path.splitext(name_file(paths[0]))[0].replace(' ',''))
    l_cmd = [] #list command to convert by pandoc

    for path in paths:
        if os.path.basename(path).endswith('.docx'):
            docx_file = dires + path
            new_name = os.path.splitext(name_file(path))[0].replace(' ','')
            container = dires + new_name
            md_file = container + '\\' + new_name.__add__(str('.md'))
            media = dires + 'media'

            os.makedirs(name=(container + "\\"), exist_ok=True)
            cmd = 'pandoc --extract-media=./ --markdown-headings=atx --wrap=none --toc -f docx -t markdown "{inp}" -o "{out}" && move "{media}" "{container}"'.format(inp=docx_file, out=md_file, media=media, container=container)

            try:
                move = 'move "{name}" "{container}"'.format(name=md_file, container=container)
                os.system(move)
            except ValueError():
                print("Error")            

            l_cmd.append(cmd)

    return l_cmd

def name_file(path):
    if os.path.basename(path).__contains__('.docx'):
        new_name = ""
        if path.__contains__("19522445_writeup-"):
            new_name = path.replace("19522445_writeup-", '')
        elif path.__contains__("19522445_"):
            new_name = path.replace("19522445_",'')
        else:
            new_name = path
        return new_name

def convert(list_path):
    #convert multi folder
    for p in list_path:
        os.chdir(p)
        #if you want to delete three first line
        delete_lines(p, 3)

        cmd = docx_2_md(p)
        for i in cmd:
            try:
                os.system(i)
            except ValueError():
                print("Error")

        print('Convert successful!')

new_folder = "convert_done"
paths = ['C:\\Users\\truon\\Desktop\\HK6\\NT213-BMW\\BT3.1\\only-docx\\', 'C:\\Users\\truon\\Desktop\\HK6\\NT213-BMW\\BT4.1\\only-docx\\', 'C:\\Users\\truon\\Desktop\\HK6\\NT213-BMW\\BT5.1\\only-docx\\']
convert(paths)
print("DONEEEEEEE!")
